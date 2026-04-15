#!/usr/bin/env python3
"""Extract {{field:type:format}} placeholders from a local .docx template.

Usage: parse_template.py <template.docx>
"""
import json
import re
import sys
from pathlib import Path

from docx import Document

PLACEHOLDER_RE = re.compile(r"\{\{\s*([^}:\s]+)(?:\s*:\s*([^:}]+))?(?:\s*:\s*([^}]+))?\s*\}\}")


def iter_text(doc):
    for p in doc.paragraphs:
        yield p.text
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    yield p.text


def main() -> int:
    if len(sys.argv) < 2:
        print("usage: parse_template.py <template.docx>", file=sys.stderr)
        return 1

    path = Path(sys.argv[1]).expanduser()
    if not path.exists():
        print(json.dumps({"error": f"file not found: {path}"}))
        return 1

    doc = Document(path)
    full_text = "\n".join(iter_text(doc))

    seen = {}
    for match in PLACEHOLDER_RE.finditer(full_text):
        key = match.group(1)
        type_spec = (match.group(2) or "text").strip()
        fmt = (match.group(3) or "").strip() or None
        seen.setdefault(key, {
            "name": key,
            "type": type_spec,
            "format": fmt,
            "required": not key.endswith("?"),
        })

    fields = sorted(seen.values(), key=lambda f: f["name"])
    print(json.dumps({
        "template": str(path),
        "total_fields": len(fields),
        "fields": fields,
    }, indent=2, ensure_ascii=False))
    return 0


if __name__ == "__main__":
    sys.exit(main())
