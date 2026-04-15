#!/usr/bin/env python3
"""Local DOCX template filler with typed placeholders.

Placeholder syntax: {{field_name:type:format}}
Types: text, text_uppercase, date, currency_number, currency_text
Format examples:
    text:##.###.###/####-##   (CNPJ mask)
    date:dd/mm/yyyy
    date:dd de MMMM de yyyy
    currency_number:##.###,##
"""
import json
import os
import re
import sys
from datetime import datetime
from pathlib import Path

from docx import Document
from num2words import num2words

MESES_PT = [
    "janeiro", "fevereiro", "março", "abril", "maio", "junho",
    "julho", "agosto", "setembro", "outubro", "novembro", "dezembro",
]

PLACEHOLDER_RE = re.compile(r"\{\{\s*([^}:\s]+)(?:\s*:\s*([^:}]+))?(?:\s*:\s*([^}]+))?\s*\}\}")


def apply_mask(value, mask):
    s = str(value)
    digits = re.sub(r"\D", "", s)
    letters = re.sub(r"[^A-Za-z]", "", s)
    out, di, li = [], 0, 0
    for ch in mask:
        if ch == "#":
            if di < len(digits):
                out.append(digits[di]); di += 1
        elif ch == "@":
            if li < len(letters):
                out.append(letters[li]); li += 1
        else:
            out.append(ch)
    return "".join(out)


def parse_date(value):
    if isinstance(value, datetime):
        return value
    for pat in ("%Y-%m-%d", "%d/%m/%Y", "%d-%m-%Y"):
        try:
            return datetime.strptime(str(value), pat)
        except ValueError:
            continue
    raise ValueError(f"unrecognized date: {value!r}")


def format_date(value, fmt):
    d = parse_date(value)
    fmt = (fmt or "dd/mm/yyyy").strip()
    if fmt == "dd/mm/yyyy":
        return d.strftime("%d/%m/%Y")
    if fmt == "dd de MMMM de yyyy":
        return f"{d.day:02d} de {MESES_PT[d.month - 1]} de {d.year}"
    return d.strftime("%d/%m/%Y")


def parse_number(value):
    if isinstance(value, (int, float)):
        return float(value)
    s = str(value).replace("R$", "").strip()
    if "," in s:
        s = s.replace(".", "").replace(",", ".")
    return float(s)


def format_currency_number(value, _mask=None):
    n = parse_number(value)
    s = f"{n:,.2f}"
    return s.replace(",", "X").replace(".", ",").replace("X", ".")


def format_currency_text(value):
    n = parse_number(value)
    return num2words(n, lang="pt_BR", to="currency")


def resolve(key, type_spec, fmt, data):
    raw = data.get(key)
    if raw is None and key.endswith("_extenso"):
        raw = data.get(key[: -len("_extenso")])
    if raw is None:
        return None
    t = (type_spec or "text").strip()
    if t == "text":
        return apply_mask(raw, fmt) if fmt else str(raw)
    if t == "text_uppercase":
        return str(raw).upper()
    if t == "date":
        return format_date(raw, fmt)
    if t == "currency_number":
        return format_currency_number(raw, fmt)
    if t == "currency_text":
        return format_currency_text(raw)
    return str(raw)


def replace_in_paragraph(paragraph, data, missing):
    if not paragraph.runs:
        return
    full = "".join(r.text for r in paragraph.runs)
    if "{{" not in full:
        return

    def sub(match):
        key, type_spec, fmt = match.group(1), match.group(2), match.group(3)
        result = resolve(key, type_spec, fmt, data)
        if result is None:
            missing.add(key)
            return match.group(0)
        return result

    new_text = PLACEHOLDER_RE.sub(sub, full)
    if new_text == full:
        return
    paragraph.runs[0].text = new_text
    for r in paragraph.runs[1:]:
        r.text = ""


def walk_paragraphs(doc):
    yield from doc.paragraphs
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                yield from cell.paragraphs


def fill_template(template_path, data, output_path):
    doc = Document(template_path)
    missing = set()
    for p in walk_paragraphs(doc):
        replace_in_paragraph(p, data, missing)
    Path(output_path).parent.mkdir(parents=True, exist_ok=True)
    doc.save(output_path)
    return missing


def main():
    if len(sys.argv) < 3:
        print("usage: fill_docx.py <template.docx> <data.json> [output.docx]")
        return 1
    template = sys.argv[1]
    data_file = sys.argv[2]
    with open(data_file) as f:
        data = json.load(f)
    if len(sys.argv) > 3:
        output = sys.argv[3]
    else:
        output = str(
            Path.home() / "Documents" / "contratos_gerados" /
            f"{Path(template).stem.replace('TEMPLATE_', '')}_filled.docx"
        )
    missing = fill_template(template, data, output)
    print(f"wrote {output}")
    if missing:
        print(f"WARNING: {len(missing)} placeholders had no data:")
        for k in sorted(missing):
            print(f"  - {k}")
    return 0


if __name__ == "__main__":
    sys.exit(main())
